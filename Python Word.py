from docx import Document
from datetime import datetime
import pandas as pd

tabela =pd.read_excel('Pasta1.xlsx')

for linha in tabela.index:
    
    documento=Document("Contrato.docx")


    nome = tabela.loc[linha,'Nome']
    item1=tabela.loc[linha,'Item1']
    item2=tabela.loc[linha,'Item2']
    item3=tabela.loc[linha,'Item3']


    referencias={
        'XXXX': nome,
        'YYYYY': item1,
        'ZZZZZ': item2,
        'WWW': item3,
        'DD':str(datetime.now().day),
        'MM':str(datetime.now().month),
        'AAAA': str(datetime.now().year)
    }


    print(documento.paragraphs)
    for paragrafo in documento.paragraphs:
        for  cod in referencias:
            valor= referencias[cod]
            paragrafo.text= paragrafo.text.replace(cod,valor)
    documento.save(f"Contrato-{nome}.docx")